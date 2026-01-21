# 파이썬 -> docx(워드파일) 생성
# pip install python-docx
from docx import Document
# 기본폰트가 영어로 지정되어 있어서 한글은 깨짐
from docx.shared import Pt

# 빈 워드문서 생성
doc = Document()

# 제목 추가하기
# 제목크기 - level 1,2,3
title = doc.add_heading("파이썬스터디", level=1)
# title에 커서를 대고 폰트를 변경
title.runs[0].font.name = "Malgun Gothic"

# 한줄추가하기
pg1 = doc.add_paragraph("파이썬 공부중")
pg1.runs[0].font.name = "Malgun Gothic"

# 파일 저장
# save("저장 경로")
doc.save("./파이썬테스트.docx")


##### os 모듈 #####
# 파일이나 폴더를 파이썬으로 다룰 때
import os

# 1. 파일이나 폴더가 존재하는지?
# os.path.exists(경로)
if os.path.exists("./파이썬테스트.docx"):
    print("파이썬테스트 파일 존재합니다!")

if not os.path.exists("./자바테스트.docx"):
    print("자바테스트 파일 존재안합니다.")

# 2. 폴더 만들기
# makedirs()
# exist_ok -> 이미 있어도 괜찮아? 네
# 이미 같은 이름의 폴더가 존재해도 에러 x
os.makedirs("내가만든폴더", exist_ok=True)

def apply_font(*texts):
    for text in texts:
        text.runs[0].font.name = "Malgun Gothic"


## 도전!) 제목: "내가만든 워드", 글: "헬로 파이썬"
# 내가만든폴더2에 저장해보세요
new_doc = Document() # 새파일

title2 = new_doc.add_heading("테스트 입니다", level=1)
# 한글폰트 지정
new_pg = new_doc.add_paragraph("헬로 파이썬222")

apply_font(title2, new_pg)
os.makedirs("내가만든폴더2", exist_ok=True)

if os.path.exists("./내가만든폴더2"):
    new_doc.save("./내가만든폴더2/연습.docx")

